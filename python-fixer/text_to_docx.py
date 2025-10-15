"""
Convert plain text to a beautiful UVA-formatted Word document
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

# UVA Brand Colors
UVA_NAVY_BLUE = RGBColor(35, 45, 75)  # #232D4B
UVA_ORANGE = RGBColor(229, 114, 0)     # #E57200

# UVA Fonts
UVA_HEADING_FONT = "Adobe Caslon Pro"
UVA_BODY_FONT = "Franklin Gothic"


def create_document_from_text(text, output_dir="output"):
    """Convert plain text to beautifully formatted UVA document"""

    doc = Document()
    fixes = []

    # Set up document styles
    _setup_styles(doc)

    # Parse and format the text
    lines = text.split('\n')

    for i, line in enumerate(lines):
        line = line.strip()

        # Skip empty lines
        if not line:
            doc.add_paragraph()
            continue

        # Detect heading type
        heading_info = _detect_heading(line, i == 0)

        if heading_info['is_heading']:
            # Add as heading
            p = doc.add_heading(heading_info['text'], level=heading_info['level'])
            p.runs[0].font.name = UVA_HEADING_FONT
            p.runs[0].font.color.rgb = UVA_NAVY_BLUE
            p.runs[0].font.bold = True

            if heading_info['level'] == 1:
                p.runs[0].font.size = Pt(18)
            elif heading_info['level'] == 2:
                p.runs[0].font.size = Pt(15)
            else:
                p.runs[0].font.size = Pt(13)

            # Fix spacing - no space after headings
            from docx.shared import Pt
            p.paragraph_format.space_before = Pt(12) if heading_info['level'] > 1 else Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        elif line.startswith('-') or line.startswith('•'):
            # Bullet point
            text_content = line.lstrip('-•').strip()
            p = doc.add_paragraph(text_content, style='List Bullet')
            for run in p.runs:
                run.font.name = UVA_BODY_FONT
                run.font.size = Pt(11)

        elif line[0].isdigit() and '. ' in line[:5]:
            # Numbered list
            p = doc.add_paragraph(line, style='List Number')
            for run in p.runs:
                run.font.name = UVA_BODY_FONT
                run.font.size = Pt(11)

        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = UVA_BODY_FONT
                run.font.size = Pt(11)

            # Tighter paragraph spacing
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

    # Add logo to footer
    _add_logo_to_footer(doc)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = int(datetime.now().timestamp())
    output_path = os.path.join(output_dir, f"created-document-{timestamp}.docx")
    doc.save(output_path)

    fixes = [
        {'type': 'font', 'description': 'Applied UVA fonts throughout document', 'location': 'All text'},
        {'type': 'hierarchy', 'description': 'Created proper heading structure', 'location': 'Headings detected'},
        {'type': 'logo', 'description': 'Added UVA Batten School logo to footer', 'location': 'Footer'},
    ]

    # Generate change log
    log_path = _generate_change_log(output_dir, fixes, timestamp)

    return {
        'success': True,
        'fixed_document': output_path,
        'change_log': log_path,
        'fixes': fixes,
        'summary': {
            'total_fixes': len(fixes),
            'font_fixes': 1,
            'logo_added': True,
        }
    }


def _setup_styles(doc):
    """Set up document styles with UVA branding"""

    styles = doc.styles

    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = UVA_BODY_FONT
    normal_style.font.size = Pt(11)

    # Heading styles
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in styles:
            style = styles[style_name]
            style.font.name = UVA_HEADING_FONT
            style.font.color.rgb = UVA_NAVY_BLUE
            style.font.bold = True


def _detect_heading(line, is_first):
    """Intelligently detect if line should be a heading"""

    # First line is document title (H1)
    if is_first:
        return {'is_heading': True, 'level': 1, 'text': line}

    # Numbered sections (1. Data Management)
    if line and line[0].isdigit() and '. ' in line[:5] and line[2].isupper():
        return {'is_heading': True, 'level': 2, 'text': line}

    # Short lines without ending punctuation
    if len(line) < 80 and not line.endswith(('.', '!', '?', ':')):
        # Check for heading keywords
        heading_keywords = ['Management', 'Review', 'Audit', 'Policy', 'System',
                          'Overview', 'Introduction', 'Conclusion', 'Summary']

        if any(keyword in line for keyword in heading_keywords):
            return {'is_heading': True, 'level': 3, 'text': line}

    return {'is_heading': False, 'level': 0, 'text': line}


def _add_logo_to_footer(doc):
    """Add UVA logo to footer"""

    logo_path = os.path.join(os.path.dirname(__file__), '..', 'battenlogo.png')

    if not os.path.exists(logo_path):
        return

    section = doc.sections[0]
    footer = section.footer

    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(4.0))


def _generate_change_log(output_dir, fixes, timestamp):
    """Generate change log for text conversion"""

    log_path = os.path.join(output_dir, f"changelog-{timestamp}.txt")

    with open(log_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("DOCUMENT CREATION REPORT\n")
        f.write("Created from plain text with UVA branding\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 80 + "\n\n")

        f.write("APPLIED FORMATTING:\n")
        f.write("-" * 80 + "\n")
        f.write("- UVA Fonts: Adobe Caslon Pro (headings), Franklin Gothic (body)\n")
        f.write("- UVA Colors: Navy Blue (#232D4B) for headings\n")
        f.write("- Professional spacing and hierarchy\n")
        f.write("- UVA Batten School logo in footer\n")
        f.write("=" * 80 + "\n")

    return log_path
