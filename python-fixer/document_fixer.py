"""
UVA Document Fixer - Python Version
Preserves original document structure while applying UVA branding
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime
from pdf2docx import Converter

# UVA Brand Colors
UVA_NAVY_BLUE = RGBColor(35, 45, 75)  # #232D4B
UVA_ORANGE = RGBColor(229, 114, 0)     # #E57200

# UVA Fonts
UVA_HEADING_FONT = "Adobe Caslon Pro"
UVA_BODY_FONT = "Franklin Gothic"

class DocumentFixer:
    def __init__(self, input_path, output_dir="output"):
        self.input_path = input_path
        self.output_dir = output_dir
        self.doc = Document(input_path)
        self.fixes = []

        # Create output directory
        os.makedirs(output_dir, exist_ok=True)

    def apply_uva_branding(self):
        """Apply UVA fonts and colors while preserving structure"""

        print(">>> apply_uva_branding() started")

        # Remove all comments first
        self._remove_comments()

        # Fix fake headings (bold/large text that should be proper headings)
        self._fix_fake_headings()

        # Fix all paragraphs
        for paragraph in self.doc.paragraphs:
            self._fix_paragraph(paragraph)

        # Fix all table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fix_paragraph(paragraph)

        # Modify styles (this affects the whole document)
        self._apply_style_changes()

        # Clean up excessive whitespace AFTER all paragraph fixes
        self._clean_whitespace()

        # Fix hyperlink sizes at XML level (python-docx doesn't handle these well)
        self._fix_hyperlinks_xml()

        # Add logo to footer
        self._add_logo_to_footer()

        return self.fixes

    def _remove_comments(self):
        """Remove all Word comments from the document"""

        try:
            # Get the document part
            document_part = self.doc.part

            # Define namespace
            w_namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

            # Remove comments from document XML
            for element in document_part.element.body:
                # Remove comment range start/end markers
                for comment_start in element.findall(f'.//{w_namespace}commentRangeStart'):
                    if comment_start.getparent() is not None:
                        comment_start.getparent().remove(comment_start)

                for comment_end in element.findall(f'.//{w_namespace}commentRangeEnd'):
                    if comment_end.getparent() is not None:
                        comment_end.getparent().remove(comment_end)

                # Remove comment references
                for comment_ref in element.findall(f'.//{w_namespace}commentReference'):
                    if comment_ref.getparent() is not None:
                        comment_ref.getparent().remove(comment_ref)

            self.fixes.append({
                'type': 'accessibility',
                'description': 'Removed all Word comments from document',
                'location': 'Entire document',
                'after': 'Comments cleaned up for final version'
            })
        except Exception as e:
            # If comment removal fails, just skip it
            print(f"Warning: Could not remove comments: {e}")
            pass

    def _fix_paragraph(self, paragraph):
        """Fix a single paragraph's formatting"""

        # Skip empty paragraphs
        if not paragraph.runs:
            return

        # Check if this is a hyperlink
        is_hyperlink = 'Hyperlink' in paragraph.style.name or 'hyperlink' in paragraph.style.name.lower()

        # Determine if it's a heading
        is_heading = paragraph.style.name.startswith('Heading') or paragraph.style.name == 'Title'

        # Determine correct font size based on style
        if is_hyperlink:
            correct_size = Pt(11)  # Hyperlinks should be body text size
        elif paragraph.style.name == 'Heading 1' or paragraph.style.name == 'Title':
            correct_size = Pt(18)
        elif paragraph.style.name == 'Heading 2':
            correct_size = Pt(15)
        elif paragraph.style.name == 'Heading 3':
            correct_size = Pt(13)
        else:
            correct_size = Pt(11)  # Body text

        for run in paragraph.runs:
            # Skip empty runs or runs without font
            if not run.text or run.font is None:
                continue

            # Debug: print paragraph style for hyperlinks
            if run.font.underline and run.font.color.rgb:
                actual_size = run.font.size.pt if run.font.size else "None"
                print(f">>> Hyperlink: style='{paragraph.style.name}', size={actual_size}pt, text='{run.text[:30]}'")

            # Apply fonts
            try:
                if is_hyperlink:
                    # For hyperlinks, only fix font size and set body font
                    run.font.name = UVA_BODY_FONT
                    # Keep hyperlink color and underline - don't change them
                elif is_heading:
                    run.font.name = UVA_HEADING_FONT
                    run.font.color.rgb = UVA_NAVY_BLUE
                    run.font.bold = True
                else:
                    run.font.name = UVA_BODY_FONT
                    # Fix colored body text - should be black (EXCEPT hyperlinks which have underline)
                    if run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
                        # Don't change color if it's a hyperlink (has underline)
                        if not run.font.underline:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            self.fixes.append({
                                'type': 'accessibility',
                                'description': 'Fixed colored body text to black',
                                'location': paragraph.text[:50]
                            })

                # For hyperlinks, set the size FIRST at XML level, THEN remove hyperlink
                has_color = run.font.color.rgb is not None and run.font.color.rgb != RGBColor(0, 0, 0)
                if run.font.underline and has_color:
                    # FIRST: Force set font size at XML level
                    rPr = run._element.rPr
                    if rPr is not None:
                        sz = rPr.get_or_add_sz()
                        sz.val = int(correct_size.pt * 2)  # Word uses half-points (11pt = 22)

                    print(f">>> Hyperlink: Set to {correct_size.pt}pt (XML): '{run.text[:30]}'")

                # THEN: Enforce standard font sizes through API
                old_size = run.font.size.pt if run.font.size else "None"
                run.font.size = correct_size


                # Remove small caps (not in UVA design standards)
                if run.font.small_caps:
                    run.font.small_caps = False
                    self.fixes.append({
                        'type': 'accessibility',
                        'description': 'Removed small caps formatting',
                        'location': paragraph.text[:50]
                    })

                # Remove all caps (accessibility issue)
                if run.font.all_caps:
                    run.font.all_caps = False
                    self.fixes.append({
                        'type': 'accessibility',
                        'description': 'Removed all caps formatting',
                        'location': paragraph.text[:50]
                    })

                # Remove underline ONLY if NOT a heading and NOT colored (hyperlinks are colored)
                has_color = run.font.color.rgb is not None and run.font.color.rgb != RGBColor(0, 0, 0)
                if run.font.underline and not run.font.underline is None:
                    # Keep underline if it's a hyperlink (has color + underline)
                    if not has_color and not is_heading:
                        run.font.underline = False
                        self.fixes.append({
                            'type': 'accessibility',
                            'description': 'Removed underline (use bold/italic for emphasis)',
                            'location': paragraph.text[:50]
                        })

                # Remove italics (not in UVA design standards for formal documents)
                if run.font.italic:
                    run.font.italic = False
                    self.fixes.append({
                        'type': 'font',
                        'description': 'Removed italic formatting',
                        'location': paragraph.text[:50]
                    })

            except Exception as e:
                # Skip runs that can't be formatted
                print(f"Warning: Could not format run: {e}")
                continue

        # Fix paragraph alignment - body text should be left-aligned
        if not is_heading and paragraph.alignment not in [None, WD_ALIGN_PARAGRAPH.LEFT]:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.fixes.append({
                'type': 'hierarchy',
                'description': 'Fixed text alignment to left',
                'location': paragraph.text[:50]
            })

        # Fix paragraph spacing
        if is_heading:
            # Headings: no space before or after (body text follows immediately)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
        else:
            # Body text: minimal space, tight spacing
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15

    def _apply_style_changes(self):
        """Modify document styles"""

        styles = self.doc.styles

        # Modify heading styles
        for i in range(1, 7):
            style_name = f'Heading {i}'
            try:
                style = styles[style_name]

                # Only set font properties if font object exists
                if style.font is not None:
                    if UVA_HEADING_FONT:
                        style.font.name = UVA_HEADING_FONT
                    if UVA_NAVY_BLUE:
                        style.font.color.rgb = UVA_NAVY_BLUE
                    style.font.bold = True

                    # Set font sizes
                    if i == 1:
                        style.font.size = Pt(18)
                    elif i == 2:
                        style.font.size = Pt(15)
                    else:
                        style.font.size = Pt(13)

                    # Set spacing - no space before or after headings
                    style.paragraph_format.space_after = Pt(0)
                    style.paragraph_format.space_before = Pt(0)

                    self.fixes.append({
                        'type': 'font',
                        'description': f'Applied UVA heading font to {style_name}',
                        'location': 'Document styles'
                    })
            except (KeyError, AttributeError) as e:
                pass

        # Modify normal style
        try:
            normal_style = styles['Normal']

            # Only set font properties if font object exists
            if normal_style.font is not None:
                if UVA_BODY_FONT:
                    normal_style.font.name = UVA_BODY_FONT
                normal_style.font.size = Pt(11)

                self.fixes.append({
                    'type': 'font',
                    'description': 'Applied UVA body font to Normal style',
                    'location': 'Document styles'
                })
        except (KeyError, AttributeError) as e:
            pass

        # Fix ALL hyperlink-related styles - should be body text size
        for style in styles:
            try:
                if 'hyperlink' in style.name.lower() or 'link' in style.name.lower():
                    if style.font is not None:
                        if UVA_BODY_FONT:
                            style.font.name = UVA_BODY_FONT
                        style.font.size = Pt(11)
                        print(f">>> Fixed style '{style.name}' to 11pt")
                        self.fixes.append({
                            'type': 'font',
                            'description': f'Set {style.name} style to body text size (11pt)',
                            'location': 'Document styles'
                        })
            except (KeyError, AttributeError) as e:
                pass

    def _clean_whitespace(self):
        """Remove excessive whitespace and empty paragraphs"""

        # Keep removing empty paragraphs until none are left
        removed_count = 0
        max_iterations = 10  # Prevent infinite loop

        print(f">>> Starting whitespace cleanup...")

        for iteration in range(max_iterations):
            paragraphs_to_remove = []

            # Find all empty or whitespace-only paragraphs
            for i, paragraph in enumerate(self.doc.paragraphs):
                # Check if paragraph has no visible text
                text = paragraph.text.strip()

                # Also check if it only contains runs with no text
                has_content = False
                if paragraph.runs:
                    for run in paragraph.runs:
                        if run.text and run.text.strip():
                            has_content = True
                            break

                if not text and not has_content:
                    paragraphs_to_remove.append(paragraph)

            print(f">>> Iteration {iteration + 1}: Found {len(paragraphs_to_remove)} empty paragraphs")

            if not paragraphs_to_remove:
                break  # No more empty paragraphs

            # Remove them
            for paragraph in paragraphs_to_remove:
                try:
                    p_element = paragraph._element
                    parent = p_element.getparent()
                    if parent is not None:
                        parent.remove(p_element)
                        removed_count += 1
                except Exception as e:
                    print(f"Warning: Could not remove empty paragraph: {e}")

        print(f">>> Finished: Removed {removed_count} empty paragraphs total")

        if removed_count > 0:
            self.fixes.append({
                'type': 'hierarchy',
                'description': f'Removed {removed_count} empty paragraphs for tight spacing',
                'location': 'Throughout document'
            })

        # Clean up multiple spaces within text and fix period spacing
        import re
        for paragraph in self.doc.paragraphs:
            # First pass: fix spacing within runs
            for run in paragraph.runs:
                if run.text:
                    original = run.text

                    # Replace multiple spaces with single space
                    cleaned = ' '.join(run.text.split())

                    # Fix period spacing: ensure exactly one space after periods
                    # This handles cases like "word.word" or "word.  word"
                    cleaned = re.sub(r'\.(?=[A-Z])', '. ', cleaned)  # Add space after period before capital letter
                    cleaned = re.sub(r'\. {2,}', '. ', cleaned)  # Replace multiple spaces after period with one

                    if original != cleaned and cleaned:
                        run.text = cleaned

            # Second pass: fix spacing between runs (especially after hyperlinks)
            runs = paragraph.runs
            for i in range(len(runs)):
                current_run = runs[i]

                # Skip empty runs
                if not current_run.text or not current_run.text.strip():
                    continue

                # Find next non-empty run
                next_run = None
                for j in range(i + 1, len(runs)):
                    if runs[j].text and runs[j].text.strip():
                        next_run = runs[j]
                        break

                if next_run:
                    # If current run ends with text (no space) and next run starts with text (no space)
                    current_ends_no_space = not current_run.text.endswith((' ', '\n', '\t'))
                    next_starts_no_space = not next_run.text.startswith((' ', '\n', '\t'))

                    # Check if we need space (current ends with letter/number, next starts with letter)
                    if current_ends_no_space and next_starts_no_space:
                        last_char = current_run.text[-1]
                        first_char = next_run.text[0]

                        # Add space if transitioning to a letter (from letter, number, or closing punctuation like ] )
                        needs_space = False
                        if first_char.isalpha():
                            if last_char.isalnum() or last_char in '])}"':
                                needs_space = True

                        if needs_space:
                            current_run.text += ' '
                            print(f">>> Added space after run: '{current_run.text[-20:]}' before '{next_run.text[:20]}'")
                            self.fixes.append({
                                'type': 'accessibility',
                                'description': 'Fixed spacing between runs',
                                'location': paragraph.text[:50]
                            })

    def _fix_hyperlinks_xml(self):
        """Fix hyperlink font sizes directly in XML (python-docx API doesn't work well for these)"""

        from lxml import etree

        print(">>> Fixing hyperlinks at XML level...")

        # Get the document XML element
        body = self.doc.element.body

        # Define XML namespaces
        nsmap = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        }

        # Find all hyperlink elements
        hyperlinks = body.findall('.//w:hyperlink', nsmap)
        fixed_count = 0
        spacing_fixes = 0

        for hyperlink in hyperlinks:
            # Check spacing before hyperlink
            parent = hyperlink.getparent()
            if parent is not None:
                children = list(parent)
                hl_index = children.index(hyperlink)

                # Check if previous element is a run and needs space before hyperlink
                if hl_index > 0:
                    prev_elem = children[hl_index - 1]
                    if prev_elem.tag == qn('w:r'):
                        prev_t = prev_elem.find('.//w:t', nsmap)
                        if prev_t is not None and prev_t.text:
                            # Check if previous text doesn't end with space
                            if not prev_t.text.endswith((' ', '\n', '\t')):
                                last_char = prev_t.text[-1]
                                # Get first char of hyperlink
                                hl_runs = hyperlink.findall('.//w:r', nsmap)
                                if hl_runs:
                                    hl_t = hl_runs[0].find('.//w:t', nsmap)
                                    if hl_t is not None and hl_t.text:
                                        first_char = hl_t.text[0]
                                        # Add space if needed
                                        if first_char.isalpha() and (last_char.isalnum() or last_char in '�'):
                                            prev_t.text += ' '
                                            # Need to preserve space
                                            if prev_t.get(qn('xml:space')) is None:
                                                prev_t.set(qn('xml:space'), 'preserve')
                                            spacing_fixes += 1
                                            print(f">>> Added space before hyperlink '{hl_t.text[:20]}' after '{prev_t.text[-20:]}'")

            # Find all runs within the hyperlink
            runs = hyperlink.findall('.//w:r', nsmap)

            for run in runs:
                # Get or create rPr (run properties)
                rPr = run.find('w:rPr', nsmap)
                if rPr is None:
                    rPr = etree.SubElement(run, qn('w:rPr'))

                # Get or create sz (font size) element
                sz = rPr.find('w:sz', nsmap)
                if sz is None:
                    sz = etree.SubElement(rPr, qn('w:sz'))

                # Set font size to 11pt (22 half-points)
                sz.set(qn('w:val'), '22')

                # Also set szCs (complex script font size)
                szCs = rPr.find('w:szCs', nsmap)
                if szCs is None:
                    szCs = etree.SubElement(rPr, qn('w:szCs'))
                szCs.set(qn('w:val'), '22')

                fixed_count += 1

                # Get text for logging
                t_elem = run.find('.//w:t', nsmap)
                if t_elem is not None and t_elem.text:
                    print(f">>> Fixed hyperlink: '{t_elem.text[:30]}' to 11pt")

        print(f">>> Fixed {fixed_count} hyperlink runs")

        if fixed_count > 0:
            self.fixes.append({
                'type': 'font',
                'description': f'Fixed {fixed_count} hyperlink font sizes to 11pt',
                'location': 'Throughout document'
            })

    def _fix_fake_headings(self):
        """Convert bold/large text paragraphs into proper heading styles"""

        print(">>> Checking for fake headings...")

        for paragraph in self.doc.paragraphs:
            # Skip if already a heading
            if paragraph.style.name.startswith('Heading') or paragraph.style.name == 'Title':
                continue

            # Skip empty paragraphs
            if not paragraph.text.strip():
                continue

            # Check if this looks like a fake heading:
            # - Short text (less than 80 chars)
            # - All runs are bold
            # - Larger font size (15pt+)

            if len(paragraph.text) < 80:
                all_bold = True
                has_hyperlink = False
                max_font_size = 0

                for run in paragraph.runs:
                    if run.text.strip():
                        if not run.font.bold:
                            all_bold = False
                        if run.font.size:
                            max_font_size = max(max_font_size, run.font.size.pt)
                        # Check if this run is a hyperlink
                        if run.font.underline and run.font.color.rgb:
                            has_hyperlink = True

                # Convert to heading based on size
                heading_style = None
                if max_font_size >= 18 and all_bold:
                    # If it's a hyperlink, make it Heading 2 (subsection)
                    # Otherwise make it Heading 1 (main section)
                    heading_style = 'Heading 2' if has_hyperlink else 'Heading 1'
                elif max_font_size >= 15 and all_bold:
                    heading_style = 'Heading 2'
                elif max_font_size >= 13 and all_bold:
                    heading_style = 'Heading 3'

                if heading_style:
                    print(f">>> Converting fake heading: '{paragraph.text[:40]}' ({max_font_size}pt bold, hyperlink={has_hyperlink}) -> {heading_style}")
                    paragraph.style = self.doc.styles[heading_style]

                    self.fixes.append({
                        'type': 'hierarchy',
                        'description': f'Converted fake heading to {heading_style}',
                        'location': paragraph.text[:50]
                    })

    def _add_logo_to_footer(self):
        """Add UVA Batten School logo to footer"""

        logo_path = os.path.join(os.path.dirname(__file__), '..', 'battenlogo.png')

        if not os.path.exists(logo_path):
            self.fixes.append({
                'type': 'logo',
                'description': 'Logo file not found, skipped',
                'location': 'Footer'
            })
            return

        # Get or create footer
        section = self.doc.sections[0]
        footer = section.footer

        # Clear existing footer
        for element in footer.paragraphs:
            element.clear()

        # Add logo
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(4.0))

        self.fixes.append({
            'type': 'logo',
            'description': 'Added UVA Batten School logo to footer',
            'location': 'Footer',
            'after': 'Logo inserted with proper branding'
        })

    def save(self, filename=None):
        """Save the fixed document"""

        if filename is None:
            timestamp = int(datetime.now().timestamp())
            filename = f"fixed-document-{timestamp}.docx"

        output_path = os.path.join(self.output_dir, filename)
        self.doc.save(output_path)

        return output_path

    def generate_change_log(self):
        """Generate a text file documenting all changes"""

        timestamp = int(datetime.now().timestamp())
        log_path = os.path.join(self.output_dir, f"changelog-{timestamp}.txt")

        with open(log_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("DOCUMENT CORRECTION REPORT\n")
            f.write("Brand Kit: UVA Batten School\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 80 + "\n\n")

            f.write("SUMMARY\n")
            f.write("-" * 80 + "\n")
            f.write(f"Total Corrections: {len(self.fixes)}\n\n")

            # Count by type
            fix_types = {}
            for fix in self.fixes:
                fix_types[fix['type']] = fix_types.get(fix['type'], 0) + 1

            for fix_type, count in fix_types.items():
                f.write(f"  {fix_type.capitalize()} Corrections: {count}\n")

            f.write("\n\nDETAILED CHANGES\n")
            f.write("-" * 80 + "\n\n")

            for i, fix in enumerate(self.fixes, 1):
                f.write(f"{i}. [{fix['type'].upper()}] {fix['description']}\n")
                f.write(f"   Location: {fix['location']}\n")
                if 'after' in fix:
                    f.write(f"   After: {fix['after']}\n")
                f.write("\n")

            f.write("\n" + "=" * 80 + "\n")
            f.write("BRAND KIT APPLIED\n")
            f.write("-" * 80 + "\n")
            f.write("Heading Font: Adobe Caslon Pro\n")
            f.write("Body Font: Franklin Gothic\n")
            f.write("Primary Color: UVA Navy Blue (#232D4B)\n")
            f.write("Accent Color: UVA Orange (#E57200)\n")
            f.write("=" * 80 + "\n")

        return log_path


def convert_pdf_to_docx(pdf_path, output_dir="output"):
    """Convert PDF to DOCX using pdf2docx"""

    print(f">>> Converting PDF to DOCX: {pdf_path}")

    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)

    # Generate output filename
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    timestamp = int(datetime.now().timestamp())
    docx_path = os.path.join(output_dir, f"{base_name}-converted-{timestamp}.docx")

    # Convert PDF to DOCX
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        print(f">>> PDF converted successfully: {docx_path}")
        return docx_path
    except Exception as e:
        print(f">>> ERROR: Failed to convert PDF: {e}")
        raise Exception(f"PDF conversion failed: {e}")


def fix_document(input_path, output_dir="output"):
    """Main function to fix a document (supports both DOCX and PDF)"""

    # Store original filename for output
    original_basename = os.path.splitext(os.path.basename(input_path))[0]

    # Check file extension
    file_ext = os.path.splitext(input_path)[1].lower()

    # If PDF, convert to DOCX first
    if file_ext == '.pdf':
        print(">>> PDF detected, converting to DOCX first...")
        docx_path = convert_pdf_to_docx(input_path, output_dir)
        input_path = docx_path
    elif file_ext != '.docx':
        raise Exception(f"Unsupported file format: {file_ext}. Only DOCX and PDF are supported.")

    fixer = DocumentFixer(input_path, output_dir)
    fixes = fixer.apply_uva_branding()

    # Save with original filename + "-Fixed"
    output_filename = f"{original_basename}-Fixed.docx"
    output_path = fixer.save(filename=output_filename)
    log_path = fixer.generate_change_log()

    return {
        'success': True,
        'fixed_document': output_path,
        'change_log': log_path,
        'fixes': fixes,
        'summary': {
            'total_fixes': len(fixes),
            'font_fixes': len([f for f in fixes if f['type'] == 'font']),
            'logo_added': any(f['type'] == 'logo' for f in fixes),
        }
    }


if __name__ == '__main__':
    # Test the fixer
    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_fixer.py <input_file.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    result = fix_document(input_file)

    print(f"✓ Document fixed: {result['fixed_document']}")
    print(f"✓ Change log: {result['change_log']}")
    print(f"✓ Total fixes: {result['summary']['total_fixes']}")
